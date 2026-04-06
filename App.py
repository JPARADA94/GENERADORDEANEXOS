import io
import math
import os
import tempfile
from typing import List, Tuple

import streamlit as st
from PIL import Image, ImageOps
from reportlab.lib.pagesizes import letter
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase.pdfmetrics import stringWidth

try:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt
    WORD_AVAILABLE = True
except Exception:
    WORD_AVAILABLE = False

st.set_page_config(
    page_title="Anexos de videoscopía",
    page_icon="🖼️",
    layout="wide",
)

# =========================================================
# CONFIGURACIÓN GENERAL
# =========================================================
PAGE_W, PAGE_H = letter
COLUMNS = 2
ROWS = 4
MAX_PER_PAGE = COLUMNS * ROWS
DEFAULT_LOGO_PATH = "/mnt/data/image.png"


def mm_to_pt(mm: float) -> float:
    return mm * 72 / 25.4


OUTER_MARGIN = mm_to_pt(12)
GAP_X = mm_to_pt(6)
GAP_Y = mm_to_pt(6)
FOOTER_H = mm_to_pt(12)
DESC_BLOCK_H = mm_to_pt(30)
TOP_SPACE = mm_to_pt(2)

USABLE_W = PAGE_W - (2 * OUTER_MARGIN)
USABLE_H = PAGE_H - (2 * OUTER_MARGIN) - FOOTER_H - DESC_BLOCK_H - TOP_SPACE
BOX_W = (USABLE_W - GAP_X) / COLUMNS
BOX_H = (USABLE_H - (GAP_Y * (ROWS - 1))) / ROWS

PDF_FONT_REGULAR = "Helvetica"
PDF_FONT_BOLD = "Helvetica-Bold"
PDF_FONT_SIZE = 12
PDF_LINE_HEIGHT = 14


# =========================================================
# UTILIDADES
# =========================================================
def registrar_fuentes_pdf():
    posibles = [
        "/usr/share/fonts/truetype/msttcorefonts/Arial.ttf",
        "/usr/share/fonts/truetype/msttcorefonts/arial.ttf",
        "/usr/share/fonts/truetype/msttcorefonts/Arial_Bold.ttf",
        "/usr/share/fonts/truetype/msttcorefonts/arialbd.ttf",
        "/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
        "/usr/share/fonts/truetype/liberation2/LiberationSans-Bold.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf",
    ]

    regular = None
    bold = None

    for p in posibles:
        if os.path.exists(p):
            low = os.path.basename(p).lower()
            if "bold" in low or "bd" in low:
                if bold is None:
                    bold = p
            else:
                if regular is None:
                    regular = p

    global PDF_FONT_REGULAR, PDF_FONT_BOLD

    try:
        if regular and "ArialCustom" not in pdfmetrics.getRegisteredFontNames():
            pdfmetrics.registerFont(TTFont("ArialCustom", regular))
            PDF_FONT_REGULAR = "ArialCustom"
        if bold and "ArialCustomBold" not in pdfmetrics.getRegisteredFontNames():
            pdfmetrics.registerFont(TTFont("ArialCustomBold", bold))
            PDF_FONT_BOLD = "ArialCustomBold"
        elif PDF_FONT_REGULAR == "ArialCustom":
            PDF_FONT_BOLD = PDF_FONT_REGULAR
    except Exception:
        PDF_FONT_REGULAR = "Helvetica"
        PDF_FONT_BOLD = "Helvetica-Bold"


registrar_fuentes_pdf()


def natural_key(name: str):
    parts = []
    current = ""
    is_digit = None

    for ch in name.lower():
        if ch.isdigit():
            if is_digit is False:
                parts.append(current)
                current = ""
            current += ch
            is_digit = True
        else:
            if is_digit is True:
                parts.append(int(current))
                current = ""
            current += ch
            is_digit = False

    if current:
        parts.append(int(current) if is_digit else current)

    return parts



def ordenar_archivos(files) -> List:
    return sorted(files, key=lambda f: natural_key(f.name))



def abrir_imagen(file) -> Image.Image:
    if hasattr(file, "seek"):
        file.seek(0)
    img = Image.open(file)
    img = ImageOps.exif_transpose(img)
    if img.mode not in ("RGB", "L"):
        img = img.convert("RGB")
    elif img.mode == "L":
        img = img.convert("RGB")
    return img



def calcular_ajuste(img_w: int, img_h: int, box_w: float, box_h: float) -> Tuple[float, float]:
    ratio = min(box_w / img_w, box_h / img_h)
    return img_w * ratio, img_h * ratio



def wrap_text_pdf(text: str, font_name: str, font_size: float, max_width: float, max_lines: int = 5) -> List[str]:
    text = (text or "").strip()
    if not text:
        return []

    words = text.split()
    lines = []
    current = ""

    for word in words:
        test = f"{current} {word}".strip()
        width = stringWidth(test, font_name, font_size)
        if width <= max_width:
            current = test
        else:
            if current:
                lines.append(current)
            current = word
            if len(lines) >= max_lines:
                break

    if current and len(lines) < max_lines:
        lines.append(current)

    return lines[:max_lines]



def draw_justified_line(pdf: canvas.Canvas, words: List[str], x: float, y: float, width: float, font_name: str, font_size: float):
    if len(words) <= 1:
        pdf.drawString(x, y, " ".join(words))
        return

    word_widths = [stringWidth(w, font_name, font_size) for w in words]
    total_words_width = sum(word_widths)
    gaps = len(words) - 1
    total_space = width - total_words_width
    if total_space <= 0:
        pdf.drawString(x, y, " ".join(words))
        return

    gap_width = total_space / gaps
    cursor_x = x
    for i, word in enumerate(words):
        pdf.drawString(cursor_x, y, word)
        cursor_x += word_widths[i]
        if i < gaps:
            cursor_x += gap_width



def draw_paragraph_continued(pdf: canvas.Canvas, label: str, text: str, x: float, y: float, total_width: float, max_lines: int = 5):
    label_width = stringWidth(label, PDF_FONT_BOLD, PDF_FONT_SIZE)
    gap = mm_to_pt(2)
    first_line_width = max(total_width - label_width - gap, mm_to_pt(20))

    words = (text or "").strip().split()
    first_line_words = []
    remaining_words = []

    current = ""
    used_count = 0
    for idx, word in enumerate(words):
        test = f"{current} {word}".strip()
        if stringWidth(test, PDF_FONT_REGULAR, PDF_FONT_SIZE) <= first_line_width:
            current = test
            used_count = idx + 1
        else:
            break

    if current:
        first_line_words = current.split()
    remaining_words = words[used_count:]

    pdf.setFont(PDF_FONT_BOLD, PDF_FONT_SIZE)
    pdf.drawString(x, y, label)

    pdf.setFont(PDF_FONT_REGULAR, PDF_FONT_SIZE)
    if first_line_words:
        draw_justified_line(
            pdf,
            first_line_words,
            x + label_width + gap,
            y,
            first_line_width,
            PDF_FONT_REGULAR,
            PDF_FONT_SIZE,
        )

    if not remaining_words:
        return

    remaining_text = " ".join(remaining_words)
    other_lines = wrap_text_pdf(
        remaining_text,
        font_name=PDF_FONT_REGULAR,
        font_size=PDF_FONT_SIZE,
        max_width=total_width,
        max_lines=max_lines - 1,
    )

    for i, line in enumerate(other_lines):
        line_y = y - ((i + 1) * PDF_LINE_HEIGHT)
        line_words = line.split()
        is_last = i == len(other_lines) - 1
        if is_last:
            pdf.drawString(x, line_y, line)
        else:
            draw_justified_line(
                pdf,
                line_words,
                x,
                line_y,
                total_width,
                PDF_FONT_REGULAR,
                PDF_FONT_SIZE,
            )



def wrap_text_word(text: str, max_chars_per_line: int = 115, max_lines: int = 5) -> str:
    text = (text or "").strip()
    if not text:
        return ""

    words = text.split()
    lines = []
    current = ""

    for word in words:
        test = f"{current} {word}".strip()
        if len(test) <= max_chars_per_line:
            current = test
        else:
            if current:
                lines.append(current)
            current = word
            if len(lines) >= max_lines:
                break

    if current and len(lines) < max_lines:
        lines.append(current)

    return "\n".join(lines[:max_lines])



def obtener_logo_fuente(logo_file):
    if logo_file is not None:
        return logo_file
    if os.path.exists(DEFAULT_LOGO_PATH):
        return DEFAULT_LOGO_PATH
    return None



def preparar_logo_pdf(logo_file, max_width=mm_to_pt(24), max_height=mm_to_pt(10)):
    src = obtener_logo_fuente(logo_file)
    if not src:
        return None, 0, 0

    img = abrir_imagen(src)
    draw_w, draw_h = calcular_ajuste(img.width, img.height, max_width, max_height)
    return ImageReader(img), draw_w, draw_h


# =========================================================
# PDF
# =========================================================
def dibujar_descripcion_pdf(pdf: canvas.Canvas, cilindro: str, descripcion: str):
    x = OUTER_MARGIN
    y = PAGE_H - OUTER_MARGIN - mm_to_pt(2)
    etiqueta = f"CILINDRO {cilindro.strip() if cilindro else 'X'}:"

    draw_paragraph_continued(
        pdf,
        etiqueta,
        descripcion,
        x,
        y,
        USABLE_W,
        max_lines=5,
    )



def dibujar_footer_pdf(pdf: canvas.Canvas, campo: str, logo_reader, logo_w: float, logo_h: float):
    text_y = OUTER_MARGIN + mm_to_pt(1)

    pdf.setFont(PDF_FONT_REGULAR, 9)
    pdf.drawString(OUTER_MARGIN, text_y, "Lubricantes Mobil")

    campo = (campo or "").strip()
    if campo:
        pdf.drawCentredString(PAGE_W / 2, text_y, campo)

    if logo_reader:
        x_logo = PAGE_W - OUTER_MARGIN - logo_w
        y_logo = OUTER_MARGIN
        pdf.drawImage(
            logo_reader,
            x_logo,
            y_logo,
            width=logo_w,
            height=logo_h,
            preserveAspectRatio=True,
            mask="auto",
        )



def generar_pdf(registros, campo: str, logo_file, cilindro: str, descripcion: str) -> bytes:
    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=letter)

    logo_reader, logo_w, logo_h = preparar_logo_pdf(logo_file)
    total_paginas = math.ceil(len(registros) / MAX_PER_PAGE)

    for page_idx in range(total_paginas):
        start = page_idx * MAX_PER_PAGE
        end = start + MAX_PER_PAGE
        lote = registros[start:end]

        dibujar_descripcion_pdf(pdf, cilindro, descripcion)

        for i, item in enumerate(lote):
            row = i // COLUMNS
            col = i % COLUMNS

            x = OUTER_MARGIN + col * (BOX_W + GAP_X)
            y_top = PAGE_H - OUTER_MARGIN - DESC_BLOCK_H - TOP_SPACE - row * (BOX_H + GAP_Y)
            y = y_top - BOX_H

            img = abrir_imagen(item["file"])
            draw_w, draw_h = calcular_ajuste(img.width, img.height, BOX_W, BOX_H)
            x_img = x + (BOX_W - draw_w) / 2
            y_img = y + (BOX_H - draw_h) / 2

            pdf.drawImage(
                ImageReader(img),
                x_img,
                y_img,
                width=draw_w,
                height=draw_h,
                preserveAspectRatio=True,
                mask="auto",
            )

        dibujar_footer_pdf(pdf, campo, logo_reader, logo_w, logo_h)
        pdf.showPage()

    pdf.save()
    buffer.seek(0)
    return buffer.getvalue()


# =========================================================
# WORD
# =========================================================
def agregar_imagen_a_word(parrafo, img: Image.Image, max_w_in=3.08, max_h_in=1.68):
    temp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
    try:
        img.save(temp.name, format="PNG")
        w_px, h_px = img.size
        ratio = min((max_w_in * 96) / w_px, (max_h_in * 96) / h_px)
        final_w = w_px * ratio / 96
        final_h = h_px * ratio / 96
        run = parrafo.add_run()
        run.add_picture(temp.name, width=Inches(final_w), height=Inches(final_h))
    finally:
        temp.close()
        if os.path.exists(temp.name):
            os.unlink(temp.name)



def agregar_logo_a_word(parrafo, logo_file, width_in=0.82):
    src = obtener_logo_fuente(logo_file)
    if not src:
        return

    img = abrir_imagen(src)
    temp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
    try:
        img.save(temp.name, format="PNG")
        parrafo.add_run().add_picture(temp.name, width=Inches(width_in))
    finally:
        temp.close()
        if os.path.exists(temp.name):
            os.unlink(temp.name)



def generar_docx(registros, campo: str, logo_file, cilindro: str, descripcion: str) -> bytes:
    if not WORD_AVAILABLE:
        return b""

    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.40)
    sec.bottom_margin = Inches(0.40)
    sec.left_margin = Inches(0.40)
    sec.right_margin = Inches(0.40)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(12)

    total_paginas = math.ceil(len(registros) / MAX_PER_PAGE)

    for page_idx in range(total_paginas):
        start = page_idx * MAX_PER_PAGE
        end = start + MAX_PER_PAGE
        lote = registros[start:end]

        p_desc = doc.add_paragraph()
        p_desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_desc.paragraph_format.space_before = Pt(0)
        p_desc.paragraph_format.space_after = Pt(4)
        p_desc.paragraph_format.line_spacing = 1.0

        r1 = p_desc.add_run(f"CILINDRO {cilindro.strip() if cilindro else 'X'}: ")
        r1.bold = True
        r1.font.name = "Arial"
        r1.font.size = Pt(12)

        texto_recortado = wrap_text_word(descripcion, max_chars_per_line=115, max_lines=5)
        r2 = p_desc.add_run(texto_recortado)
        r2.font.name = "Arial"
        r2.font.size = Pt(12)

        tabla = doc.add_table(rows=ROWS, cols=COLUMNS)
        tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
        tabla.autofit = False

        for row in tabla.rows:
            row.height = Inches(2.05)
            for cell in row.cells:
                cell.width = Inches(3.75)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cell.text = ""

        for i, item in enumerate(lote):
            r = i // COLUMNS
            c = i % COLUMNS
            cell = tabla.cell(r, c)

            p_img = cell.paragraphs[0]
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(0)
            p_img.paragraph_format.space_after = Pt(0)
            agregar_imagen_a_word(p_img, abrir_imagen(item["file"]))

        p_gap = doc.add_paragraph()
        p_gap.paragraph_format.space_before = Pt(1)
        p_gap.paragraph_format.space_after = Pt(0)

        footer = doc.add_table(rows=1, cols=3)
        footer.alignment = WD_TABLE_ALIGNMENT.CENTER
        footer.autofit = False

        footer.rows[0].cells[0].width = Inches(2.2)
        footer.rows[0].cells[1].width = Inches(3.3)
        footer.rows[0].cells[2].width = Inches(1.3)

        c0, c1, c2 = footer.rows[0].cells

        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p0.paragraph_format.space_after = Pt(0)
        r0 = p0.add_run("Lubricantes Mobil")
        r0.font.name = "Arial"
        r0.font.size = Pt(9)

        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.space_after = Pt(0)
        rm = p1.add_run((campo or "").strip())
        rm.font.name = "Arial"
        rm.font.size = Pt(9)

        p2 = c2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p2.paragraph_format.space_after = Pt(0)
        agregar_logo_a_word(p2, logo_file)

        if page_idx < total_paginas - 1:
            doc.add_section(WD_SECTION.NEW_PAGE)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# =========================================================
# ESTADO Y ORDEN MANUAL
# =========================================================
def inicializar_registros(files):
    ordenados = ordenar_archivos(files)
    regs = []
    for idx, f in enumerate(ordenados, start=1):
        regs.append({
            "uid": f"{idx}_{f.name}_{getattr(f, 'size', 0)}",
            "file": f,
        })
    return regs



def mover_arriba(registros, idx):
    if idx > 0:
        registros[idx - 1], registros[idx] = registros[idx], registros[idx - 1]



def mover_abajo(registros, idx):
    if idx < len(registros) - 1:
        registros[idx + 1], registros[idx] = registros[idx], registros[idx + 1]


# =========================================================
# INTERFAZ
# =========================================================
st.title("🖼️ Generador de anexos de videoscopía")

with st.container(border=True):
    st.markdown("#### Datos generales")
    campo = st.text_input("Campo", placeholder="Ej: Estación SANTS")
    cilindro = st.text_input("Cilindro", placeholder="Ej: 1L, 2R")
    descripcion = st.text_area("Descripción de hallazgos", height=140)
    logo_file = st.file_uploader("Logo opcional", type=["png", "jpg", "jpeg"], key="logo_uploader")

    if logo_file is None and os.path.exists(DEFAULT_LOGO_PATH):
        st.caption("Se usará el logo predeterminado de Mobil.")

uploaded_files = st.file_uploader(
    "Sube imágenes",
    type=["jpg", "jpeg", "png"],
    accept_multiple_files=True,
    key="images_uploader",
)

if uploaded_files:
    current_ids = [
        f"{i}_{f.name}_{getattr(f, 'size', 0)}"
        for i, f in enumerate(ordenar_archivos(uploaded_files), start=1)
    ]

    if "registros_imagenes" not in st.session_state:
        st.session_state.registros_imagenes = inicializar_registros(uploaded_files)
    else:
        existing_ids = [r["uid"] for r in st.session_state.registros_imagenes]
        if existing_ids != current_ids:
            st.session_state.registros_imagenes = inicializar_registros(uploaded_files)

    registros = st.session_state.registros_imagenes

    st.success(f"Se cargaron {len(registros)} imagen(es).")

    with st.container(border=True):
        st.markdown("#### Reordenar imágenes")
        st.caption("Usa los botones para cambiar el orden final del anexo.")

        for idx, item in enumerate(registros):
            c1, c2 = st.columns([3, 1])

            with c1:
                st.image(abrir_imagen(item["file"]), use_container_width=True)
                st.caption(item["file"].name)

            with c2:
                st.write("")
                if st.button("⬆️ Subir", key=f"up_{item['uid']}", use_container_width=True):
                    mover_arriba(registros, idx)
                    st.rerun()
                if st.button("⬇️ Bajar", key=f"down_{item['uid']}", use_container_width=True):
                    mover_abajo(registros, idx)
                    st.rerun()

    with st.expander("Ver orden final"):
        for i, item in enumerate(registros, start=1):
            st.write(f"{i}. {item['file'].name}")

    total_paginas = math.ceil(len(registros) / MAX_PER_PAGE)
    st.info(f"Se generarán {total_paginas} página(s), con máximo {MAX_PER_PAGE} imágenes por página.")

    pdf_bytes = generar_pdf(registros, campo, logo_file, cilindro, descripcion)

    if WORD_AVAILABLE:
        word_bytes = generar_docx(registros, campo, logo_file, cilindro, descripcion)

        col1, col2 = st.columns(2)
        with col1:
            st.download_button(
                "📄 Descargar PDF",
                data=pdf_bytes,
                file_name="anexos_videoscopia.pdf",
                mime="application/pdf",
                use_container_width=True,
            )
        with col2:
            st.download_button(
                "📝 Descargar Word",
                data=word_bytes,
                file_name="anexos_videoscopia.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
    else:
        st.download_button(
            "📄 Descargar PDF",
            data=pdf_bytes,
            file_name="anexos_videoscopia.pdf",
            mime="application/pdf",
            use_container_width=True,
        )
        st.warning("La exportación a Word no está disponible porque falta instalar python-docx.")
else:
    st.info("Sube imágenes para continuar.")

with st.container(border=True):
    st.markdown("#### requirements.txt")
    st.code(
        """streamlit
pillow
reportlab
python-docx""",
        language="text",
    )

    st.markdown("#### Ejecución local")
    st.code("streamlit run App.py", language="bash")

